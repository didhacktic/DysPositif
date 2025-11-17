#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de diagnostic pour les zones de texte VML dans DysPositif.
À exécuter avec : python3 diagnostic_vml.py chemin/vers/fichier.docx
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import zipfile
from lxml import etree

def diagnostic_vml(filepath):
    """Analyse complète d'un document pour les zones de texte VML."""
    
    print("=" * 80)
    print("DIAGNOSTIC VML - DysPositif")
    print("=" * 80)
    print(f"\nFichier: {filepath}")
    
    if not os.path.exists(filepath):
        print(f"✗ ERREUR: Le fichier n'existe pas!")
        return False
    
    # 1. Analyse XML brute
    print("\n[1] Analyse de la structure XML...")
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            doc_xml = zf.read('word/document.xml')
        
        root = etree.fromstring(doc_xml)
        
        # Compter les éléments VML
        txbx_count = len(list(root.iter(qn('w:txbxContent'))))
        pict_count = len(list(root.iter(qn('w:pict'))))
        
        VML_NS = 'urn:schemas-microsoft-com:vml'
        vshape_count = len(list(root.iter('{%s}shape' % VML_NS)))
        vtextbox_count = len(list(root.iter('{%s}textbox' % VML_NS)))
        
        print(f"    w:txbxContent: {txbx_count}")
        print(f"    w:pict: {pict_count}")
        print(f"    v:shape: {vshape_count}")
        print(f"    v:textbox: {vtextbox_count}")
        
        if txbx_count == 0:
            print("    ⚠ Aucune zone de texte VML détectée dans le XML!")
        else:
            print(f"    ✓ {txbx_count} zone(s) de texte VML trouvée(s)")
        
    except Exception as e:
        print(f"    ✗ Erreur lors de l'analyse XML: {e}")
        return False
    
    # 2. Test avec python-docx
    print("\n[2] Test avec python-docx...")
    try:
        doc = Document(filepath)
        print(f"    ✓ Document chargé")
        print(f"    Paragraphes normaux: {len(doc.paragraphs)}")
        print(f"    Tableaux: {len(doc.tables)}")
        print(f"    Inline shapes: {len(doc.inline_shapes)}")
    except Exception as e:
        print(f"    ✗ Erreur: {e}")
        return False
    
    # 3. Test de la fonction _iter_vml_textbox_paragraphs
    print("\n[3] Test de détection VML (_iter_vml_textbox_paragraphs)...")
    try:
        # Fonction de détection VML
        def _iter_vml_textbox_paragraphs(doc):
            root = doc.element
            for txbx_content in root.iter(qn('w:txbxContent')):
                for p_elem in txbx_content.iter(qn('w:p')):
                    try:
                        yield Paragraph(p_elem, txbx_content)
                    except Exception:
                        pass
        
        vml_paras = list(_iter_vml_textbox_paragraphs(doc))
        print(f"    Paragraphes VML détectés: {len(vml_paras)}")
        
        if len(vml_paras) == 0:
            print("    ✗ PROBLÈME: Aucun paragraphe VML détecté!")
            print("    → La fonction de détection ne trouve pas les zones de texte.")
        else:
            print("    ✓ Détection réussie!")
            for i, p in enumerate(vml_paras, 1):
                text = p.text[:80] + "..." if len(p.text) > 80 else p.text
                print(f"      {i}. '{text}'")
                print(f"         Runs: {len(p.runs)}")
        
    except Exception as e:
        print(f"    ✗ Erreur: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # 4. Test de formatage
    print("\n[4] Test d'application de formatage...")
    try:
        from docx.shared import Pt
        
        if len(vml_paras) == 0:
            print("    ⊘ Pas de paragraphes VML à formater")
        else:
            # Appliquer un formatage de test
            for p in vml_paras:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(16)
            
            print(f"    ✓ Formatage appliqué à {len(vml_paras)} paragraphe(s)")
            
            # Sauvegarder pour test
            test_output = filepath.replace('.docx', '_TEST_VML.docx')
            doc.save(test_output)
            print(f"    ✓ Document de test sauvegardé: {test_output}")
            
            # Recharger et vérifier
            doc2 = Document(test_output)
            vml_paras2 = list(_iter_vml_textbox_paragraphs(doc2))
            
            if len(vml_paras2) > 0:
                success = True
                for p in vml_paras2:
                    for run in p.runs:
                        if run.font.name != "Arial":
                            success = False
                            print(f"    ✗ Police non appliquée: {run.font.name}")
                
                if success:
                    print("    ✓✓✓ Le formatage persiste après sauvegarde!")
                else:
                    print("    ✗ Le formatage n'a pas persisté")
            else:
                print("    ✗ Les paragraphes VML ont disparu après sauvegarde!")
        
    except Exception as e:
        print(f"    ✗ Erreur: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    # Résumé
    print("\n" + "=" * 80)
    print("RÉSUMÉ")
    print("=" * 80)
    
    if txbx_count > 0 and len(vml_paras) > 0:
        print("✓ Les zones de texte VML sont détectées correctement")
        print("✓ Le formatage devrait fonctionner avec DysPositif")
        print(f"\nVérifiez le fichier de test: {test_output}")
        print("Si le texte dans les zones de texte est formaté en Arial 16pt,")
        print("alors le code fonctionne correctement!")
        return True
    elif txbx_count > 0 and len(vml_paras) == 0:
        print("✗ PROBLÈME: Les zones VML existent mais ne sont pas détectées")
        print("→ Il y a un problème avec la fonction de détection")
        return False
    else:
        print("ℹ Ce document ne contient pas de zones de texte VML")
        print("→ Les zones de texte sont peut-être d'un autre type (DrawingML)")
        return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 diagnostic_vml.py chemin/vers/fichier.docx")
        print("\nExemple:")
        print("  python3 diagnostic_vml.py ~/didhacktic/DysPositif/Fichiers_test/06_Texte_image_zt.docx")
        sys.exit(1)
    
    filepath = sys.argv[1]
    success = diagnostic_vml(filepath)
    sys.exit(0 if success else 1)
