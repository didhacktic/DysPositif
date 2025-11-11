# core/syllables_and_mute.py
import re
from docx.shared import RGBColor

# Couleurs
ROUGE = RGBColor(220, 20, 60)
BLEU = RGBColor(30, 144, 255)
GRIS = RGBColor(180, 180, 180)

# Muettes
MUETTES = {
    'ent': ['vent', 'dent', 'cent', 'lent', 'gent', 'absent', 'présent'],
    'ez': ['chez', 'rez', 'nez'],
    'ait': ['fait', 'trait', 'plait'],
}

# Pattern mot (inclut apostrophes, tirets)
WORD_PATTERN = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ'’-]+")

# Normalisation accents – clés de longueur 1
ACCENT_MAP = str.maketrans({
    'à': 'a', 'á': 'a', 'â': 'a', 'ã': 'a', 'ä': 'a', 'å': 'a',
    'è': 'e', 'é': 'e', 'ê': 'e', 'ë': 'e',
    'ì': 'i', 'í': 'i', 'î': 'i', 'ï': 'i',
    'ò': 'o', 'ó': 'o', 'ô': 'o', 'õ': 'o', 'ö': 'o',
    'ù': 'u', 'ú': 'u', 'û': 'u', 'ü': 'u',
    'ç': 'c', 'ñ': 'n', 'ý': 'y', 'ÿ': 'y',
    'æ': 'ae', 'œ': 'oe',
    'À': 'A', 'Á': 'A', 'Â': 'A', 'Ã': 'A', 'Ä': 'A', 'Å': 'A',
    'È': 'E', 'É': 'E', 'Ê': 'E', 'Ë': 'E',
    'Ì': 'I', 'Í': 'I', 'Î': 'I', 'Ï': 'I',
    'Ò': 'O', 'Ó': 'O', 'Ô': 'O', 'Õ': 'O', 'Ö': 'O',
    'Ù': 'U', 'Ú': 'U', 'Û': 'U', 'Ü': 'U',
    'Ç': 'C', 'Ñ': 'N', 'Ý': 'Y', 'Ÿ': 'Y',
    'Æ': 'AE', 'Œ': 'OE'
})

def normalize(text: str) -> str:
    return text.lower().translate(ACCENT_MAP)

def est_muette(mot_lower: str) -> str:
    for suffixe, exceptions in MUETTES.items():
        if mot_lower.endswith(suffixe) and mot_lower not in exceptions:
            return suffixe
    return None

def split_syllables(word_norm: str):
    vowels = "aeiouy"
    parts = []
    current = ""
    for c in word_norm:
        current += c
        if c in vowels:
            parts.append(current)
            current = ""
    if current:
        if parts:
            parts[-1] += current
        else:
            parts.append(current)
    return parts or [word_norm]

def apply_syllables_and_mute(doc):
    counter = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue

        # Sauvegarder style du paragraphe
        style = paragraph.style
        alignment = paragraph.alignment

        # Extraire tout le texte
        full_text = paragraph.text
        paragraph.clear()

        i = 0
        while i < len(full_text):
            char = full_text[i]

            # Ponctuation / espaces
            if not WORD_PATTERN.match(char):
                run = paragraph.add_run(char)
                i += 1
                continue

            # Trouver mot complet
            match = WORD_PATTERN.search(full_text[i:])
            if not match:
                run = paragraph.add_run(char)
                i += 1
                continue

            word = match.group()
            word_start = i
            i += len(word)

            word_lower = word.lower()
            muette_suffix = est_muette(word_lower)
            is_muette = muette_suffix and word_lower.endswith(muette_suffix)

            # Découpage syllabique
            word_norm = normalize(word)
            syll_parts = split_syllables(word_norm)

            pos = 0
            for part in syll_parts:
                part_len = len(part)
                orig_part = word[pos:pos + part_len]

                # Ajustement si décalage accent
                if normalize(orig_part) != part:
                    orig_part = word[pos:pos + 1]
                    pos += 1
                else:
                    pos += part_len

                if not orig_part:
                    continue

                run = paragraph.add_run(orig_part)

                # Couleur syllabe
                if is_muette and pos == len(word):
                    run.font.color.rgb = GRIS
                else:
                    run.font.color.rgb = ROUGE if counter % 2 == 0 else BLEU
                    counter += 1

        # Restaurer style paragraphe
        paragraph.style = style
        if alignment:
            paragraph.alignment = alignment