# -------------------------------------------------
# core/numbers_multicolor.py – Chiffres multicolores (PRÉSERVE TOUT)
# -------------------------------------------------
import re
from docx.shared import RGBColor

COULEURS = {
    '0': RGBColor(100,100,100), '1': RGBColor(200,200,200),
    '2': RGBColor(220,20,60),   '3': RGBColor(144,238,144),
    '4': RGBColor(148,0,211),   '5': RGBColor(255,215,0),
    '6': RGBColor(0,0,139),     '7': RGBColor(0,0,0),
    '8': RGBColor(139,69,19),   '9': RGBColor(173,216,230),
}

def apply_multicolor_numbers(doc):
    containers = [doc]
    for table in doc.tables:
        containers.extend(table._cells)
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame') and shape.text_frame:
            containers.append(shape.text_frame)

    for container in containers:
        for p in container.paragraphs:
            if not p.text.strip():
                continue
            texte = p.text
            runs_data = [(r.text, r.bold, r.italic, r.underline, r.font.color.rgb if r.font.color and r.font.color.rgb else None) for r in p.runs]
            p.clear()
            i = 0
            run_idx = 0
            while i < len(texte):
                c = texte[i]
                if c.isdigit():
                    r = p.add_run(c)
                    r.font.color.rgb = COULEURS.get(c, RGBColor(0,0,0))
                    # Restaurer style original
                    if run_idx < len(runs_data) and runs_data[run_idx][0]:
                        r.bold, r.italic, r.underline = runs_data[run_idx][1:4]
                        if runs_data[run_idx][4]:
                            r.font.color.rgb = runs_data[run_idx][4]  # priorité style original
                        runs_data[run_idx] = (runs_data[run_idx][0][1:],) + runs_data[run_idx][1:]
                        if not runs_data[run_idx][0]:
                            run_idx += 1
                else:
                    r = p.add_run(c)
                    if run_idx < len(runs_data) and runs_data[run_idx][0]:
                        r.bold, r.italic, r.underline = runs_data[run_idx][1:4]
                        if runs_data[run_idx][4]:
                            r.font.color.rgb = runs_data[run_idx][4]
                        runs_data[run_idx] = (runs_data[run_idx][0][1:],) + runs_data[run_idx][1:]
                        if not runs_data[run_idx][0]:
                            run_idx += 1
                i += 1