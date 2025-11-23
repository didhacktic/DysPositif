#!/usr/bin/env python3
# -*- coding: utf-8 -*-
#
# Syllabation via pylirecouleur (Arkaline, GPL v3)
# Source originale : https://framagit.org/arkaline/pylirecouleur
# Copié et adapté pour DysPositif (GPL v3)
#
# Modifications apportées :
# - Ajout de l'import sys avant manipulation de sys.path (corrige NameError potentiel)
# - Consolidation de l'import de la fonction `syllables` pour éviter la redondance
# - Clarification : WORD_PATTERN ne contient PAS le tiret (-). Le tiret est traité
#   comme séparateur afin d'être cohérent avec le traitement des noms composés
#   (ex. "Jean-Michel") et éviter des recollages indésirables par d'autres modules.
#

import sys
import re
import os
from docx.shared import RGBColor

# Ajout du chemin src du dépôt pylirecouleur (si présent dans l'arborescence du projet)
# On insère ce chemin au début de sys.path afin que l'import suivant trouve le module.
pylire_root = os.path.join(os.path.dirname(__file__), '..', 'pylirecouleur')
if os.path.exists(pylire_root):
    src_path = os.path.join(pylire_root, 'src')
    if src_path not in sys.path:
        sys.path.insert(0, src_path)

# Import direct de la fonction de segmentation syllabique depuis lirecouleur.
# Important : cet import doit intervenir après la modification éventuelle de sys.path.
from lirecouleur.word import syllables  # Fonction principale de segmentation

# Couleurs utilisées pour l'alternance syllabique
COUL_SYLL = [RGBColor(220, 20, 60), RGBColor(30, 144, 255)]

# Pattern acceptant lettres latines, lettres accentuées françaises, apostrophes typographiques, etc.
# NOTE : le tiret '-' a été retiré de la classe afin que les mots liés par un tiret
# soient traités comme mot + séparateur + mot (cohérent avec la logique de parsing char-by-char).
WORD_PATTERN = re.compile(r"[A-Za-z\u00C0-\u00D6\u00D8-\u00F6\u00F8-\u00FF'\u2019]+")

# Table de normalisation pour supprimer/mapper les accents avant la segmentation
ACCENT_MAP = str.maketrans({
    '\u00e0': 'a', '\u00e1': 'a', '\u00e2': 'a', '\u00e3': 'a', '\u00e4': 'a', '\u00e5': 'a',
    '\u00e8': 'e', '\u00e9': 'e', '\u00ea': 'e', '\u00eb': 'e',
    '\u00ec': 'i', '\u00ed': 'i', '\u00ee': 'i', '\u00ef': 'i',
    '\u00f2': 'o', '\u00f3': 'o', '\u00f4': 'o', '\u00f5': 'o', '\u00f6': 'o',
    '\u00f9': 'u', '\u00fa': 'u', '\u00fb': 'u', '\u00fc': 'u',
    '\u00e7': 'c', '\u0153': 'oe', '\u00e6': 'ae',
    '\u00c0': 'A', '\u00c1': 'A', '\u00c2': 'A', '\u00c3': 'A', '\u00c4': 'A', '\u00c5': 'A',
    '\u00c8': 'E', '\u00c9': 'E', '\u00ca': 'E', '\u00cb': 'E',
    '\u00cc': 'I', '\u00cd': 'I', '\u00ce': 'I', '\u00cf': 'I',
    '\u00d2': 'O', '\u00d3': 'O', '\u00d4': 'O', '\u00d5': 'O', '\u00d6': 'O',
    '\u00d9': 'U', '\u00da': 'U', '\u00db': 'U', '\u00dc': 'U',
    '\u00c7': 'C', '\u00d1': 'N', '\u00dd': 'Y', '\u0178': 'Y',
    '\u00c6': 'AE', '\u0152': 'OE'
})

def normalize(word: str) -> str:
    """Retourne une version normalisée (minuscule, accents mappés) du mot."""
    return word.lower().translate(ACCENT_MAP)

def apply_syllables(doc):
    """
    Parcourt le document (paragraphes, cellules de tableau, zones texte inline, zones de texte VML)
    et remplace chaque mot par des runs colorés par syllabe.

    Méthode :
    - On extrait le texte du paragraphe, on vide le paragraphe (p.clear())
    - On parcourt caractère par caractère en segmentant les mots avec WORD_PATTERN
    - Pour chaque mot, on normalise et on appelle la fonction syllables(mot_norm)
      fournie par lirecouleur, puis on recrée des runs colorés en s'alignant
      sur la graphie originale (heuristique pour gérer les accents).
    """
    from core.utils import _iter_textbox_paragraphs
    
    counter = [0]  # compteur partagé pour alterner les couleurs
    containers = [doc]

    # Inclure les cellules de tableaux
    for table in doc.tables:
        containers.extend(table._cells)

    # Inclure les zones de texte inline (s'il y en a)
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame') and shape.text_frame:
            containers.append(shape.text_frame)
    
    # Traiter d'abord les containers standards
    all_paragraphs = []
    for container in containers:
        all_paragraphs.extend(container.paragraphs)
    
    # Ajouter les paragraphes des zones de texte VML
    all_paragraphs.extend(_iter_textbox_paragraphs(doc))

    # Séparateurs (hors espaces) considérés 'à part' => on les rend tels quels
    # NOTE: Les espaces ne sont PAS dans SEPARATORS pour permettre leur attachement aux mots
    SEPARATORS = set("\t\n\r.,;:!?()[]{}\u00ab\u00bb\u201c\u201d''/\\-–—*+=<>@#$%^&~")

    for p in all_paragraphs:
        if not p.text.strip():
            continue
        texte = p.text
        
        # Préserver le format du paragraphe avant de le vider
        # Sauvegarder les propriétés importantes pour éviter les problèmes de mise en page
        try:
            # Sauvegarder l'alignement, l'espacement, etc.
            para_format = p.paragraph_format
            keep_together = para_format.keep_together
            keep_with_next = para_format.keep_with_next
            widow_control = para_format.widow_control
        except:
            keep_together = None
            keep_with_next = None
            widow_control = None
        
        p.clear()
        
        # Restaurer les propriétés de paragraphe pour éviter les ruptures
        try:
            if keep_together is not None:
                p.paragraph_format.keep_together = True  # Forcer à garder ensemble
            if keep_with_next is not None:
                p.paragraph_format.keep_with_next = False  # Ne pas forcer avec le suivant
            if widow_control is not None:
                p.paragraph_format.widow_control = True
            # Désactiver les sauts de page automatiques
            p.paragraph_format.page_break_before = False
        except:
            pass
        i = 0
        while i < len(texte):
            c = texte[i]
            
            # Traiter les retours à la ligne et tabulations
            if c in '\n\r\t':
                p.add_run(c)
                i += 1
                continue
            
            # Traiter les séparateurs (hors espaces)
            if c in SEPARATORS:
                # Collecter les séparateurs consécutifs
                sep_start = i
                while i < len(texte) and texte[i] in SEPARATORS:
                    i += 1
                p.add_run(texte[sep_start:i])
                continue
            
            # Traiter les espaces seuls (pas suivis d'un mot)
            if c == ' ':
                # Regarder si suivi d'un mot ou d'autres espaces
                space_start = i
                while i < len(texte) and texte[i] == ' ':
                    i += 1
                # Si pas de mot après les espaces, les ajouter seuls
                if i >= len(texte) or not WORD_PATTERN.match(texte[i:]):
                    p.add_run(texte[space_start:i])
                    continue
                # Sinon, les espaces seront traités avec le mot suivant
                spaces_before = texte[space_start:i]
            else:
                spaces_before = ''

            match = WORD_PATTERN.match(texte[i:])
            if match:
                mot = match.group()
                mot_norm = normalize(mot)

                # Segmentation via lirecouleur
                try:
                    syll_parts = syllables(mot_norm)
                except Exception:
                    syll_parts = [mot]  # fallback: mot entier

                # Collecter les espaces APRÈS le mot pour les attacher au dernier run
                next_pos = i + len(mot)
                spaces_after = ''
                while next_pos < len(texte) and texte[next_pos] == ' ':
                    spaces_after += texte[next_pos]
                    next_pos += 1

                # Ajouter d'abord les espaces avant (si collectés précédemment)
                if spaces_before:
                    # Attacher aux espaces le même run que la première syllabe
                    # Pour éviter une rupture, on les ajoute au début du mot
                    pass  # Seront ajoutés avec la première syllabe

                # Recomposer les syllabes avec la graphie d'origine
                pos = 0
                for si, syl in enumerate(syll_parts):
                    target = syl
                    L = len(target)
                    part = mot[pos:pos+L] if pos+L <= len(mot) else mot[pos:]
                    
                    # Ajouter les espaces avant avec la première syllabe
                    if si == 0 and spaces_before:
                        part = spaces_before + part
                    
                    # Ajouter les espaces après avec la dernière syllabe
                    if si == len(syll_parts) - 1 and spaces_after:
                        part = part + spaces_after
                    
                    # Créer run coloré
                    run = p.add_run(part)
                    color = COUL_SYLL[counter[0] % len(COUL_SYLL)]
                    try:
                        run.font.color.rgb = color
                    except Exception:
                        pass
                    counter[0] += 1
                    pos += len(mot[pos:pos+L]) if pos+L <= len(mot) else len(mot[pos:])

                # Si reste des caractères (rare)
                if pos < len(mot):
                    rest = mot[pos:]
                    if spaces_after and si == len(syll_parts) - 1:
                        rest = rest + spaces_after
                    run = p.add_run(rest)
                    try:
                        run.font.color.rgb = COUL_SYLL[counter[0] % len(COUL_SYLL)]
                    except Exception:
                        pass
                    counter[0] += 1

                # Avancer i pour inclure le mot ET les espaces après
                i = next_pos
                spaces_before = ''  # Réinitialiser
            else:
                # caractère isolé (non matched by WORD_PATTERN) => ajouter tel quel
                p.add_run(texte[i])
                i += 1