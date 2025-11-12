#!/usr/bin/env python3
# -*- coding: UTF-8 -*-
"""
core/syllables_mute.py
Pipeline combiné : applique la coloration syllabique puis le grisage des lettres muettes
en utilisant la même stratégie de fichier temporaire que l'implémentation précédente.
Expose apply_syllables_mute(doc, input_filepath) — retourne le Document résultant
(avec les syllabes colorées et les muettes grisées).
"""

import tempfile
import os
from docx import Document
from typing import Optional

from ui.interface import update_progress
from .syllables import apply_syllables
from .mute_letters import apply_mute_letters


def apply_syllables_mute(doc: Document, input_filepath: str) -> Document:
    """
    Applique la coloration syllabique puis le grisage des lettres muettes sur `doc`.
    Stratégie :
      1) Appliquer apply_syllables(doc)
      2) Sauvegarder doc dans un fichier temporaire
      3) Charger temp_doc = Document(temp_path) et appliquer apply_mute_letters(temp_doc)
      4) Sauvegarder temp_doc et renvoyer temp_doc (caller doit utiliser le Document retourné)
      5) Supprimer le fichier temporaire

    Retourne le Document modifié (temp_doc). Lance les exceptions en cas d'erreur.
    """
    # Étape 1 : coloration syllabique
    update_progress(50, "Étape 1/2 : Coloration syllabique...")
    try:
        apply_syllables(doc)
    except Exception:
        update_progress(0, "Échec : coloration syllabique")
        raise

    # Sauvegarde temporaire du document syllabé
    fd, temp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(temp_path)
    except Exception:
        try:
            os.unlink(temp_path)
        except Exception:
            pass
        raise

    # Étape 2 : appliquer muettes sur la version syllabée
    update_progress(75, "Étape 2/2 : Grisage des lettres muettes...")
    try:
        temp_doc = Document(temp_path)
        apply_mute_letters(temp_doc)
        temp_doc.save(temp_path)
    except Exception:
        try:
            os.unlink(temp_path)
        except Exception:
            pass
        update_progress(0, "Échec : application des muettes")
        raise

    # Charger et retourner le document résultant (temp_doc)
    try:
        result_doc = Document(temp_path)
    finally:
        # Nettoyage du fichier temporaire (on peut le supprimer après avoir chargé result_doc)
        try:
            os.unlink(temp_path)
        except Exception:
            pass

    update_progress(95, "Traitement syllabes+muettes terminé")
    return result_doc