# core/utils.py
"""
Utilitaires sûrs pour modifications fines sur docx (python-docx).

Fonctions principales exposées :
- apply_font_consistently(doc, police_name, taille_pt_value, include_tables=False, include_headers_footers=True)
    -> applique nom et taille de police sur le corps, tableaux (optionnel), headers/footers (optionnel)
    -> ne modifie PAS run.font.spacing (tracking/kerning)

- apply_spacing_and_line_spacing(doc, espacement: bool, interlignes: bool)
    -> wrapper : n'applique pas le tracking, mais applique l'interlignage si demandé

- split_run_and_color(run, start, end, color_hex) et safe_color_substring_in_paragraph(...)
    -> opérations sûres pour colorer une sous-portion d'un run
"""
from __future__ import annotations
from copy import deepcopy
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

# -----------------------
# Helpers généraux
# -----------------------
def _iter_body_paragraphs(doc: Document):
    """Itère uniquement les paragraphes du corps principal (exclut headers/footers/tables)."""
    try:
        body_element = doc.element.body
        from docx.text.paragraph import Paragraph
        for child in body_element:
            if child.tag.endswith('}p'):
                parent = child.getparent()
                if parent is not None and parent.tag == body_element.tag:
                    try:
                        yield Paragraph(child, doc.element.body)
                    except Exception:
                        pass
    except Exception:
        # Fallback
        for p in doc.paragraphs:
            yield p

def _iter_textbox_paragraphs(doc: Document):
    """Itère les paragraphes situés dans les zones de texte (w:txbxContent)."""
    from docx.text.paragraph import Paragraph
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    try:
        for txbx in doc.element.body.iter():
            if txbx.tag.endswith('}txbxContent'):
                for p_elem in txbx.iter():
                    if p_elem.tag == f'{w_ns}p':
                        try:
                            yield Paragraph(p_elem, doc)
                        except Exception:
                            pass
    except Exception:
        pass

def _iter_table_paragraphs(doc: Document):
    """Itère les paragraphes situés dans les cellules de tableaux."""
    for table in getattr(doc, "tables", ()):
        for cell in getattr(table, "_cells", ()):
            for p in getattr(cell, "paragraphs", ()):
                yield p

# -----------------------
# Fonctions principales
# -----------------------
def apply_font_consistently(doc: Document, police_name: str, taille_pt_value: int,
                            include_tables: bool = False,
                            include_headers_footers: bool = True):
    """Applique nom et taille de police de manière cohérente.

    Portée :
      - Corps principal (paragraphes hors tableaux)
      - Tableaux du corps (si include_tables=True)
      - Headers / footers (paragraphes + tableaux) si include_headers_footers=True

    Ne modifie pas le tracking (run.font.spacing).
    """
    taille = Pt(taille_pt_value)

    # Corps principal
    for p in _iter_body_paragraphs(doc):
        for run in p.runs:
            try:
                run.font.name = police_name
                run.font.size = taille
            except Exception:
                pass

    # Tableaux corps
    if include_tables:
        for p in _iter_table_paragraphs(doc):
            for run in p.runs:
                try:
                    run.font.name = police_name
                    run.font.size = taille
                except Exception:
                    pass

    # Headers / Footers
    if include_headers_footers:
        try:
            for section in doc.sections:
                for container in (section.header, section.footer):
                    # Paragraphes directs
                    for p in getattr(container, 'paragraphs', ()):
                        for run in p.runs:
                            try:
                                run.font.name = police_name
                                run.font.size = taille
                            except Exception:
                                pass
                    # Tableaux dans headers/footers
                    for table in getattr(container, 'tables', ()):
                        for row in getattr(table, 'rows', ()):
                            for cell in getattr(row, 'cells', ()):
                                for p in getattr(cell, 'paragraphs', ()):
                                    for run in p.runs:
                                        try:
                                            run.font.name = police_name
                                            run.font.size = taille
                                        except Exception:
                                            pass
        except Exception:
            pass

def apply_line_spacing(doc: Document, interlignes_value: Optional[float] = None,
                       include_headers_footers: bool = True, 
                       include_tables: bool = True):
    """Applique un interlignage (multiplicateur) sur paragraphes."""
    if interlignes_value is None:
        return
    
    # Corps principal
    for p in _iter_body_paragraphs(doc):
        try:
            p.paragraph_format.line_spacing = interlignes_value
        except Exception:
            pass
    
    # Tableaux du corps
    if include_tables:
        for p in _iter_table_paragraphs(doc):
            try:
                p.paragraph_format.line_spacing = interlignes_value
            except Exception:
                pass
    
    # Headers / Footers
    if include_headers_footers:
        try:
            for section in doc.sections:
                for container in (section.header, section.footer):
                    for p in getattr(container, 'paragraphs', ()):
                        try:
                            p.paragraph_format.line_spacing = interlignes_value
                        except Exception:
                            pass
                    for table in getattr(container, 'tables', ()):
                        for row in getattr(table, 'rows', ()):
                            for cell in getattr(row, 'cells', ()):
                                for p in getattr(cell, 'paragraphs', ()):
                                    try:
                                        p.paragraph_format.line_spacing = interlignes_value
                                    except Exception:
                                        pass
        except Exception:
            pass

def apply_spacing_and_line_spacing(doc: Document, espacement: bool, interlignes: bool):
    """
    Wrapper de compatibilité pour l'ancienne API.
    - espacement : bool => IGNORÉ (pour éviter espacement excessif)
    - interlignes : bool => applique interlignage 1.5 si True, 1.0 si False
    """
    if interlignes:
        apply_line_spacing(doc, 1.5)
        compress_double_empty_lines(doc)
    else:
        # Forcer minimum 1.0 pour éviter texte trop serré
        apply_line_spacing(doc, 1.0)


def compress_double_empty_lines(doc: Document,
                                include_headers_footers: bool = True,
                                include_tables: bool = True):
    """Réduit toute séquence de paragraphes vides consécutifs à UNE seule ligne vide."""
    
    def _is_blank(paragraph) -> bool:
        try:
            return (paragraph.text or '').strip() == ''
        except Exception:
            return False

    def _compress_paragraph_list(paragraphs):
        prev_blank = False
        for p in list(paragraphs):
            if _is_blank(p):
                if prev_blank:
                    try:
                        el = p._p
                        parent = el.getparent()
                        parent.remove(el)
                    except Exception:
                        pass
                else:
                    prev_blank = True
            else:
                prev_blank = False

    # Corps principal
    try:
        body_paragraphs = []
        body_element = doc.element.body
        Paragraph = getattr(docx.text.paragraph, 'Paragraph', None)
        if Paragraph:
            for child in body_element:
                if child.tag.endswith('}p'):
                    body_paragraphs.append(Paragraph(child, body_element))
            _compress_paragraph_list(body_paragraphs)
    except Exception:
        pass

    # Tableaux
    if include_tables:
        try:
            for table in doc.tables:
                for cell in table._cells:
                    _compress_paragraph_list(cell.paragraphs)
        except Exception:
            pass

    # Headers / Footers
    if include_headers_footers:
        try:
            for section in doc.sections:
                for container in (section.header, section.footer):
                    _compress_paragraph_list(getattr(container, 'paragraphs', ()))
                    if include_tables:
                        for table in getattr(container, 'tables', ()):
                            for row in getattr(table, 'rows', ()):
                                for cell in getattr(row, 'cells', ()):
                                    _compress_paragraph_list(getattr(cell, 'paragraphs', ()))
        except Exception:
            pass

# -----------------------
# Manipulation sûre de runs (split / clone / color)
# -----------------------
def _clone_run_element_with_text(run, text: str):
    """Clone l'élément XML d'un run en conservant les propriétés et remplace le texte."""
    new_r = deepcopy(run._r)
    for t in new_r.iterfind('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        t.text = text
    return new_r

def _insert_run_after(run, new_r_element):
    """Insère new_r_element (élément XML) juste après run._r"""
    run._r.addnext(new_r_element)

def _set_color_on_run_element(run_element, color_hex: str):
    """Ajoute/modifie la couleur (w:color) sur un élément run (XML)."""
    rpr = run_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    if rpr is None:
        rpr = OxmlElement('w:rPr')
        run_element.insert(0, rpr)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), color_hex)
    # supprimer éventuelles couleurs existantes
    for child in list(rpr):
        if child.tag.endswith('}color'):
            rpr.remove(child)
    rpr.append(color_el)

def split_run_and_color(run, start: int, end: int, color_hex: str) -> None:
    """
    Découpe le run en head|middle|tail et applique color_hex uniquement sur middle.
    Préserve les propriétés rPr du run original (deepcopy).
    start inclusif, end exclusif (indices Python).
    """
    text = run.text or ""
    if not text:
        return
    if start < 0:
        start = 0
    if end > len(text):
        end = len(text)
    if start >= end:
        return

    head = text[:start]
    middle = text[start:end]
    tail = text[end:]

    try:
        original_r = deepcopy(run._r)
    except Exception:
        original_r = deepcopy(run._r)

    run.text = head

    middle_el = deepcopy(original_r)
    for t in middle_el.iterfind('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        t.text = middle
    tail_el = deepcopy(original_r)
    for t in tail_el.iterfind('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        t.text = tail

    _insert_run_after(run, middle_el)
    middle_el.addnext(tail_el)

    _set_color_on_run_element(middle_el, color_hex)

    if tail == "":
        try:
            tail_el.getparent().remove(tail_el)
        except Exception:
            pass

def safe_color_substring_in_paragraph(paragraph, substring: str, color_hex: str, first_only: bool = False):
    """
    Colorer la première (ou toutes) occurrences d'une sous-chaîne contenue DANS UN SEUL run.
    Limitation: n'interprète pas les occurrences traversant plusieurs runs.
    """
    if not substring:
        return
    for run in paragraph.runs:
        t = run.text or ""
        idx = t.find(substring)
        if idx >= 0:
            split_run_and_color(run, idx, idx + len(substring), color_hex)
            if first_only:
                return

# -----------------------
# Formatage zones de texte (w:txbxContent dans VML shapes)
# -----------------------
def apply_textbox_formatting(doc: Document, police_name: str, taille_pt_value: int, interlignes_value: Optional[float] = None):
    """
    Applique la police et augmente légèrement l'interlignage sur les zones de texte (w:txbxContent).
    Conserve la taille d'origine de chaque zone de texte.
    
    Les colorations syllabiques sont déjà appliquées par apply_syllables qui inclut
    maintenant les zones de texte. Cette fonction applique la police et un léger
    espacement supplémentaire (15% de plus que l'original).
    
    À appeler APRÈS apply_font_consistently et les colorations.
    """
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    v_ns = 'urn:schemas-microsoft-com:vml'
    import re
    
    try:
        # Formatage du contenu (police + interlignage)
        # Note: la hauteur est ajustée en post-traitement via _fix_textbox_heights()
        for txbx in doc.element.body.iter():
            if not txbx.tag.endswith('}txbxContent'):
                continue
            
            # Appliquer interlignage aux paragraphes
            for p in txbx.iter():
                if p.tag == f'{w_ns}p':
                    pPr = p.find(f'{w_ns}pPr')
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p.insert(0, pPr)
                    
                    spacing = pPr.find(f'{w_ns}spacing')
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    
                    # Appliquer un interlignage uniforme et régulier
                    # On utilise 1.05× (252 twips) - très proche de simple, juste un petit ajustement
                    # lineRule="auto" garantit un espacement uniforme entre toutes les lignes
                    spacing.set(qn('w:line'), '252')
                    spacing.set(qn('w:lineRule'), 'auto')
                    
                    # Minimal espacement avant
                    spacing.set(qn('w:before'), '20')
                    spacing.set(qn('w:after'), '0')
            
            # Appliquer la police aux runs
            for r in txbx.iter():
                if r.tag == f'{w_ns}r':
                    rPr = r.find(f'{w_ns}rPr')
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)
                    
                    # Appliquer UNIQUEMENT la police (ne touche pas sz/szCs)
                    rFonts = rPr.find(f'{w_ns}rFonts')
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:ascii'), police_name)
                    rFonts.set(qn('w:hAnsi'), police_name)
    except Exception:
        # Ne pas faire échouer le traitement si le formatage des zones de texte échoue
        pass
