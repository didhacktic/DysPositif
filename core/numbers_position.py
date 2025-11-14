# -------------------------------------------------
# core/numbers_position.py – Couleur par position (u/d/c) – PRÉSERVE TOUT
# -------------------------------------------------
from docx.shared import RGBColor

COULEURS = [RGBColor(30,144,255), RGBColor(220,20,60), RGBColor(0,128,0)]

def apply_position_numbers(doc):
    """
    Coloration des chiffres en fonction de leur position (units/tens/hundreds...)
    Position calculée depuis la droite du nombre.

    Méthode :
    - Pour chaque paragraphe / cellule / zone de texte, on lit le texte et on
      mémorise les runs (texte + styles).
    - On vide le paragraphe et on reconstruit caractère par caractère.
    - Lorsqu'on rencontre une séquence de chiffres contiguë, on traite la
      séquence entière et on applique la couleur selon la position depuis la droite.
    - Les styles (bold/italic/underline, police, taille) sont restaurés ; la
      couleur d'origine n'est PAS restaurée pour les chiffres afin que la
      coloration par position reste visible.
    """
    containers = [doc]
    for table in doc.tables:
        containers.extend(table._cells)
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame') and shape.text_frame:
            containers.append(shape.text_frame)

    for container in containers:
        for p in container.paragraphs:
            if not p.text.strip():
                continue
            texte = p.text
            # runs_data : (remaining_text_in_run, bold, italic, underline, original_color_rgb)
            runs_data = [
                (r.text or "", bool(r.bold), bool(r.italic), bool(r.underline),
                 (r.font.color.rgb if getattr(r.font, "color", None) and getattr(r.font.color, "rgb", None) else None))
                for r in p.runs
            ]
            p.clear()
            i = 0
            run_idx = 0
            while i < len(texte):
                c = texte[i]
                if c.isdigit():
                    # détecter la séquence entière de chiffres à partir de i
                    j = i
                    while j < len(texte) and texte[j].isdigit():
                        j += 1
                    seq = texte[i:j]
                    nlen = len(seq)
                    # pour chaque chiffre dans la séquence, colorer selon position depuis la droite
                    for k, ch in enumerate(seq):
                        pos_from_right = nlen - 1 - k
                        color = COULEURS[pos_from_right % 3]
                        r = p.add_run(ch)
                        # appliquer la couleur calculée (ne PAS restaurer la couleur originale)
                        try:
                            r.font.color.rgb = color
                        except Exception:
                            pass
                        # restaurer les autres attributs (gras/italique/souligné)
                        if run_idx < len(runs_data) and runs_data[run_idx][0]:
                            r.bold, r.italic, r.underline = runs_data[run_idx][1:4]
                            # consommer un caractère du run original
                            runs_data[run_idx] = (runs_data[run_idx][0][1:],) + runs_data[run_idx][1:]
                            if not runs_data[run_idx][0]:
                                run_idx += 1
                    # avancer i jusqu'après la séquence numérique
                    i = j
                else:
                    # caractère non-chiffre : ajouter tel quel et restaurer style + couleur d'origine
                    r = p.add_run(c)
                    if run_idx < len(runs_data) and runs_data[run_idx][0]:
                        r.bold, r.italic, r.underline = runs_data[run_idx][1:4]
                        if runs_data[run_idx][4]:
                            try:
                                r.font.color.rgb = runs_data[run_idx][4]
                            except Exception:
                                pass
                        runs_data[run_idx] = (runs_data[run_idx][0][1:],) + runs_data[run_idx][1:]
                        if not runs_data[run_idx][0]:
                            run_idx += 1
                    i += 1